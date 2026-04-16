import os, json, uuid, hashlib, io, tempfile, base64, logging, threading
from datetime import datetime
from pathlib import Path
from flask import (Flask, render_template, request, redirect, url_for,
                   session, jsonify, send_file, send_from_directory, abort)
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 PageBreak, Image as RLImage, Table,
                                 TableStyle, HRFlowable)
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from PIL import Image as PILImage

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'estudio-criativo-secret-2024')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100 MB

BASE = Path(__file__).parent

def _ensure_dir(path: Path) -> Path:
    """Create directory, falling back to a local path if permission is denied."""
    try:
        path.mkdir(parents=True, exist_ok=True)
        return path
    except PermissionError:
        fallback = BASE / 'data'
        logging.warning(f'[storage] cannot create {path}, falling back to {fallback}')
        fallback.mkdir(parents=True, exist_ok=True)
        return fallback

# In production, set DATA_DIR env var to a persistent disk path (e.g. /data)
_data_dir = _ensure_dir(Path(os.environ.get('DATA_DIR', str(BASE / 'data'))))
DATA_FILE = _data_dir / 'projects.json'
_uploads_target = Path(os.environ.get('UPLOADS_DIR', str(BASE / 'static' / 'uploads')))
UPLOADS_DIR = _ensure_dir(_uploads_target)

# ── Auth ──────────────────────────────────────────────────────────────────────
ADMIN_PASSWORD_HASH = hashlib.sha256(
    os.environ.get('ADMIN_PASSWORD', 'estudio2024').encode()
).hexdigest()

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'heic', 'tif', 'tiff', 'bmp'}

# ── Google Drive backup ───────────────────────────────────────────────────────
# Set these env vars in Render dashboard:
#   GDRIVE_SERVICE_ACCOUNT_B64  → base64-encoded service account JSON
#   GDRIVE_FOLDER_ID            → ID of the Drive folder shared with the service account
_gdrive_service = None
GDRIVE_FOLDER_ID = os.environ.get('GDRIVE_FOLDER_ID', '')

def _init_gdrive():
    global _gdrive_service
    if _gdrive_service is not None:
        return _gdrive_service
    b64 = os.environ.get('GDRIVE_SERVICE_ACCOUNT_B64', '')
    if not b64 or not GDRIVE_FOLDER_ID:
        return None
    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
        creds_json = json.loads(base64.b64decode(b64).decode('utf-8'))
        creds = service_account.Credentials.from_service_account_info(
            creds_json,
            scopes=['https://www.googleapis.com/auth/drive.file']
        )
        _gdrive_service = build('drive', 'v3', credentials=creds, cache_discovery=False)
        return _gdrive_service
    except Exception as e:
        logging.warning(f'[gdrive] init failed: {e}')
        return None

def _gdrive_file_id(service, name):
    """Return the Drive file ID for `name` inside GDRIVE_FOLDER_ID, or None."""
    try:
        q = f"name='{name}' and '{GDRIVE_FOLDER_ID}' in parents and trashed=false"
        res = service.files().list(q=q, fields='files(id)').execute()
        files = res.get('files', [])
        return files[0]['id'] if files else None
    except Exception:
        return None

def _backup_file_to_drive(local_path: Path, drive_name: str):
    """Upload or update a file on Drive. Runs in background thread."""
    def _run():
        service = _init_gdrive()
        if not service:
            return
        try:
            from googleapiclient.http import MediaFileUpload
            mime = 'application/json' if str(local_path).endswith('.json') else 'application/octet-stream'
            media = MediaFileUpload(str(local_path), mimetype=mime, resumable=False)
            existing_id = _gdrive_file_id(service, drive_name)
            if existing_id:
                service.files().update(fileId=existing_id, media_body=media).execute()
            else:
                meta = {'name': drive_name, 'parents': [GDRIVE_FOLDER_ID]}
                service.files().create(body=meta, media_body=media, fields='id').execute()
            logging.info(f'[gdrive] backed up {drive_name}')
        except Exception as e:
            logging.warning(f'[gdrive] backup failed for {drive_name}: {e}')
    threading.Thread(target=_run, daemon=True).start()

def _restore_from_drive():
    """On startup: if projects.json is missing, try to pull it from Drive."""
    if DATA_FILE.exists():
        return
    service = _init_gdrive()
    if not service:
        return
    try:
        from googleapiclient.http import MediaIoBaseDownload
        file_id = _gdrive_file_id(service, 'projects.json')
        if not file_id:
            return
        request_dl = service.files().get_media(fileId=file_id)
        buf = io.BytesIO()
        dl = MediaIoBaseDownload(buf, request_dl)
        done = False
        while not done:
            _, done = dl.next_chunk()
        DATA_FILE.parent.mkdir(parents=True, exist_ok=True)
        DATA_FILE.write_bytes(buf.getvalue())
        logging.info('[gdrive] restored projects.json from Drive')
    except Exception as e:
        logging.warning(f'[gdrive] restore failed: {e}')

# ── Data helpers ──────────────────────────────────────────────────────────────
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_data():
    if DATA_FILE.exists():
        with open(DATA_FILE, encoding='utf-8') as f:
            return json.load(f)
    return {"projects": []}

def save_data(data):
    DATA_FILE.parent.mkdir(parents=True, exist_ok=True)
    tmp = DATA_FILE.with_suffix('.tmp')
    with open(tmp, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    tmp.replace(DATA_FILE)
    # async backup to Google Drive — never blocks the request
    _backup_file_to_drive(DATA_FILE, 'projects.json')

def find_page(data, pid, sid, pgid):
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: return None
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: return None
    return next((pg for pg in section['pages'] if pg['id'] == pgid), None)

def require_auth(f):
    from functools import wraps
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get('authenticated'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return wrapper

# ── Auth routes ───────────────────────────────────────────────────────────────
@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        pw = request.form.get('password', '')
        if hashlib.sha256(pw.encode()).hexdigest() == ADMIN_PASSWORD_HASH:
            session['authenticated'] = True
            return redirect(url_for('index'))
        error = 'senha incorreta'
    return render_template('login.html', error=error)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

# ── Main app ──────────────────────────────────────────────────────────────────
@app.route('/')
@require_auth
def index():
    data = load_data()
    return render_template('index.html', projects=data['projects'])

# ── Notes API ─────────────────────────────────────────────────────────────────
@app.route('/api/notes', methods=['GET'])
@require_auth
def get_notes():
    data = load_data()
    return jsonify({'content': data.get('notes', '')})

@app.route('/api/notes', methods=['PUT'])
@require_auth
def save_notes():
    data = load_data()
    data['notes'] = request.json.get('content', '')
    save_data(data)
    return jsonify({'ok': True})

# ── Project API ───────────────────────────────────────────────────────────────
@app.route('/api/projects', methods=['GET'])
@require_auth
def get_projects():
    return jsonify(load_data())

@app.route('/api/projects', methods=['POST'])
@require_auth
def create_project():
    data = load_data()
    body = request.json
    project = {
        'id': str(uuid.uuid4()),
        'name': body['name'],
        'type': body['type'],
        'created': datetime.now().isoformat(),
        'sections': build_sections(body['type'])
    }
    data['projects'].append(project)
    save_data(data)
    return jsonify(project)

@app.route('/api/projects/<pid>', methods=['DELETE'])
@require_auth
def delete_project(pid):
    data = load_data()
    data['projects'] = [p for p in data['projects'] if p['id'] != pid]
    save_data(data)
    return jsonify({'ok': True})

@app.route('/api/projects/<pid>/sections/<sid>/pages', methods=['POST'])
@require_auth
def add_page(pid, sid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: abort(404)
    body = request.json
    page = {
        'id': str(uuid.uuid4()),
        'name': body['name'],
        'title': body['name'],
        'content': '',
        'scenes': [],
        'images': [],
        'board_cards': [],
        'board_connections': [],
        'form_answers': {}
    }
    section['pages'].append(page)
    save_data(data)
    return jsonify(page)

@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>', methods=['PUT'])
@require_auth
def update_page(pid, sid, pgid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: abort(404)
    page = next((pg for pg in section['pages'] if pg['id'] == pgid), None)
    if not page: abort(404)
    body = request.json
    allowed = {'title', 'content', 'scenes', 'images', 'board_cards', 'board_connections', 'form_answers'}
    for k, v in body.items():
        if k in allowed:
            page[k] = v
    save_data(data)
    return jsonify(page)

@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>', methods=['DELETE'])
@require_auth
def delete_page(pid, sid, pgid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: abort(404)
    section['pages'] = [pg for pg in section['pages'] if pg['id'] != pgid]
    save_data(data)
    return jsonify({'ok': True})

# ── Scenes API ────────────────────────────────────────────────────────────────
@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>/scenes', methods=['POST'])
@require_auth
def add_scene(pid, sid, pgid):
    data = load_data()
    page = find_page(data, pid, sid, pgid)
    if not page: abort(404)
    body = request.json
    if 'scenes' not in page: page['scenes'] = []
    scene = {
        'id': str(uuid.uuid4()),
        'title': body.get('title', 'nova cena'),
        'content': body.get('content', ''),
        'notes': body.get('notes', '')
    }
    page['scenes'].append(scene)
    save_data(data)
    return jsonify(scene)

@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>/scenes/<scid>', methods=['PUT'])
@require_auth
def update_scene(pid, sid, pgid, scid):
    data = load_data()
    page = find_page(data, pid, sid, pgid)
    if not page: abort(404)
    scene = next((s for s in page.get('scenes', []) if s['id'] == scid), None)
    if not scene: abort(404)
    body = request.json
    for k in ('title', 'content', 'notes'):
        if k in body: scene[k] = body[k]
    save_data(data)
    return jsonify(scene)

@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>/scenes/<scid>', methods=['DELETE'])
@require_auth
def delete_scene(pid, sid, pgid, scid):
    data = load_data()
    page = find_page(data, pid, sid, pgid)
    if not page: abort(404)
    page['scenes'] = [s for s in page.get('scenes', []) if s['id'] != scid]
    save_data(data)
    return jsonify({'ok': True})

@app.route('/api/projects/<pid>/sections/<sid>/move-scene', methods=['POST'])
@require_auth
def move_scene(pid, sid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: abort(404)
    body = request.json
    from_page = next((p for p in section['pages'] if p['id'] == body['from_page']), None)
    to_page   = next((p for p in section['pages'] if p['id'] == body['to_page']), None)
    if not from_page or not to_page: abort(404)
    scene = next((s for s in from_page.get('scenes', []) if s['id'] == body['scene_id']), None)
    if not scene: abort(404)
    from_page['scenes'] = [s for s in from_page['scenes'] if s['id'] != body['scene_id']]
    if 'scenes' not in to_page: to_page['scenes'] = []
    idx = min(body.get('to_index', len(to_page['scenes'])), len(to_page['scenes']))
    to_page['scenes'].insert(idx, scene)
    save_data(data)
    return jsonify({'ok': True})

# ── Upload ────────────────────────────────────────────────────────────────────
@app.route('/api/upload', methods=['POST'])
@require_auth
def upload_image():
    if 'file' not in request.files:
        return jsonify({'error': 'no file'}), 400
    f = request.files['file']
    if not f.filename or not allowed_file(f.filename):
        return jsonify({'error': 'invalid file'}), 400
    ext = f.filename.rsplit('.', 1)[1].lower()
    filename = f'{uuid.uuid4().hex}.{ext}'
    path = UPLOADS_DIR / filename
    f.save(path)
    # create web-viewable thumbnail (TIFF/BMP → JPEG)
    try:
        img = PILImage.open(path)
        img.thumbnail((800, 800))
        if ext in ('tif', 'tiff', 'bmp', 'heic'):
            thumb_name = f'thumb_{uuid.uuid4().hex}.jpg'
            img.convert('RGB').save(UPLOADS_DIR / thumb_name, 'JPEG', quality=92)
        else:
            thumb_name = f'thumb_{filename}'
            img.save(UPLOADS_DIR / thumb_name, quality=92)
    except Exception:
        thumb_name = filename
    # async backup of original + thumb to Drive
    _backup_file_to_drive(UPLOADS_DIR / filename, filename)
    if thumb_name != filename:
        _backup_file_to_drive(UPLOADS_DIR / thumb_name, thumb_name)
    return jsonify({'url': f'/uploads/{filename}', 'thumb': f'/uploads/{thumb_name}', 'name': f.filename})

@app.route('/uploads/<path:filename>')
def serve_upload(filename):
    return send_from_directory(UPLOADS_DIR, filename)

@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>/images', methods=['POST'])
@require_auth
def attach_image(pid, sid, pgid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: abort(404)
    page = next((pg for pg in section['pages'] if pg['id'] == pgid), None)
    if not page: abort(404)
    body = request.json
    if 'images' not in page:
        page['images'] = []
    img = {'id': str(uuid.uuid4()), 'url': body['url'], 'thumb': body.get('thumb', body['url']), 'caption': body.get('caption', ''), 'name': body.get('name', '')}
    page['images'].append(img)
    save_data(data)
    return jsonify(img)

@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>/images/<imgid>', methods=['DELETE'])
@require_auth
def detach_image(pid, sid, pgid, imgid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: abort(404)
    page = next((pg for pg in section['pages'] if pg['id'] == pgid), None)
    if not page: abort(404)
    page['images'] = [i for i in page.get('images', []) if i['id'] != imgid]
    save_data(data)
    return jsonify({'ok': True})

# ── Board API ─────────────────────────────────────────────────────────────────
@app.route('/api/projects/<pid>/sections/<sid>/pages/<pgid>/board', methods=['PUT'])
@require_auth
def update_board(pid, sid, pgid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    section = next((s for s in project['sections'] if s['id'] == sid), None)
    if not section: abort(404)
    page = next((pg for pg in section['pages'] if pg['id'] == pgid), None)
    if not page: abort(404)
    body = request.json
    page['board_cards'] = body.get('cards', [])
    page['board_connections'] = body.get('connections', [])
    save_data(data)
    return jsonify({'ok': True})

# ── PDF Export ────────────────────────────────────────────────────────────────
@app.route('/api/projects/<pid>/export', methods=['GET'])
@require_auth
def export_pdf(pid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    styles = getSampleStyleSheet()
    _type_hex = {'doc': '#1A2744', 'book': '#2B5231', 'photo': '#A31C2C', 'blog': '#5C6470'}
    TYPE_COLOR = colors.HexColor(_type_hex.get(project['type'], '#1A2744'))

    s_cover_title = ParagraphStyle('CoverTitle', fontSize=32, leading=38, textColor=TYPE_COLOR, spaceAfter=8, fontName='Helvetica-Bold')
    s_cover_sub = ParagraphStyle('CoverSub', fontSize=14, leading=20, textColor=colors.HexColor('#888780'), spaceAfter=4)
    s_cover_meta = ParagraphStyle('CoverMeta', fontSize=11, leading=16, textColor=colors.HexColor('#B4B2A9'))
    s_section = ParagraphStyle('Section', fontSize=18, leading=24, textColor=TYPE_COLOR, spaceAfter=6, spaceBefore=20, fontName='Helvetica-Bold')
    s_page_title = ParagraphStyle('PageTitle', fontSize=14, leading=20, textColor=colors.HexColor('#2C2C2A'), spaceAfter=8, spaceBefore=12, fontName='Helvetica-Bold')
    s_body = ParagraphStyle('Body', fontSize=11, leading=17, textColor=colors.HexColor('#444441'), spaceAfter=6)
    s_caption = ParagraphStyle('Caption', fontSize=9, leading=13, textColor=colors.HexColor('#888780'), spaceAfter=8, alignment=TA_CENTER)
    s_card_title = ParagraphStyle('CardTitle', fontSize=11, leading=15, textColor=colors.HexColor('#2C2C2A'), fontName='Helvetica-Bold')
    s_card_body = ParagraphStyle('CardBody', fontSize=10, leading=14, textColor=colors.HexColor('#5F5E5A'))

    story = []

    # Cover
    _pt_map = {'doc': 'Documentário', 'book': 'Livro', 'photo': 'Fotografia', 'blog': 'Blog / Não-Ficção'}
    pt = _pt_map.get(project['type'], project['type'])
    created = project.get('created', '')[:10] if project.get('created') else ''
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph(project['name'], s_cover_title))
    story.append(Paragraph(pt, s_cover_sub))
    story.append(Paragraph(f'exportado em {datetime.now().strftime("%d/%m/%Y")}  ·  criado em {created}', s_cover_meta))
    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width='100%', thickness=2, color=TYPE_COLOR))
    story.append(PageBreak())

    # Sections
    for section in project.get('sections', []):
        if not section.get('pages'): continue
        story.append(Paragraph(f'{section["icon"]}  {section["label"].upper()}', s_section))
        story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#D3D1C7')))
        story.append(Spacer(1, 0.3*cm))

        for page in section.get('pages', []):
            story.append(Paragraph(page.get('title') or page.get('name', ''), s_page_title))

            if page.get('content'):
                for line in page['content'].split('\n'):
                    line = line.strip()
                    if not line: story.append(Spacer(1, 0.2*cm)); continue
                    if line.startswith('## '):
                        story.append(Paragraph(line[3:], ParagraphStyle('H2', fontSize=13, leading=18, fontName='Helvetica-Bold', textColor=colors.HexColor('#3C3489'), spaceAfter=4, spaceBefore=8)))
                    elif line.startswith('- '):
                        story.append(Paragraph(f'• {line[2:]}', ParagraphStyle('Li', fontSize=11, leading=16, leftIndent=14, textColor=colors.HexColor('#444441'), spaceAfter=2)))
                    elif line.startswith('> '):
                        story.append(Paragraph(line[2:], ParagraphStyle('Quote', fontSize=11, leading=16, leftIndent=20, textColor=colors.HexColor('#888780'), spaceAfter=4, borderPad=4)))
                    else:
                        story.append(Paragraph(line, s_body))

            # Images grid
            imgs = page.get('images', [])
            if imgs:
                story.append(Spacer(1, 0.3*cm))
                img_row = []
                for img_data in imgs:
                    img_path = BASE / img_data['url'].lstrip('/')
                    if img_path.exists():
                        try:
                            rl_img = RLImage(str(img_path), width=6*cm, height=4.5*cm)
                            caption = img_data.get('caption') or img_data.get('name', '')
                            cell = [rl_img, Paragraph(caption, s_caption)]
                            img_row.append(cell)
                        except Exception:
                            pass
                if img_row:
                    cols = min(len(img_row), 3)
                    rows_data = [img_row[i:i+cols] for i in range(0, len(img_row), cols)]
                    for row in rows_data:
                        while len(row) < cols:
                            row.append([''])
                        t = Table(row, colWidths=[6.5*cm]*cols)
                        t.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('TOPPADDING', (0,0), (-1,-1), 4), ('BOTTOMPADDING', (0,0), (-1,-1), 4)]))
                        story.append(t)
                    story.append(Spacer(1, 0.3*cm))

            # Board cards
            cards = page.get('board_cards', [])
            if cards:
                story.append(Spacer(1, 0.3*cm))
                story.append(Paragraph('storyboard', ParagraphStyle('BoardLabel', fontSize=9, fontName='Helvetica-Bold', textColor=colors.HexColor('#888780'), spaceAfter=6, letterSpacing=1)))
                card_table_data = []
                row = []
                for i, card in enumerate(cards):
                    act_idx = card.get('act', 0)
                    act_colors = ['#E4E8F2','#E2EFE4','#F5E4E6','#FBF2E0']
                    bg = act_colors[act_idx % len(act_colors)]
                    cell_content = [
                        Paragraph(f'#{str(card.get("id",""))[:4]}  {card.get("tag","")}', ParagraphStyle('CardNum', fontSize=8, textColor=colors.HexColor('#888780'), spaceAfter=2)),
                        Paragraph(card.get('title',''), s_card_title),
                    ]
                    if card.get('notes'):
                        cell_content.append(Paragraph(card['notes'], s_card_body))
                    if card.get('duration'):
                        cell_content.append(Paragraph(card['duration'], ParagraphStyle('Dur', fontSize=8, textColor=colors.HexColor('#B4B2A9'), spaceBefore=4)))
                    row.append(cell_content)
                    if len(row) == 3 or i == len(cards)-1:
                        while len(row) < 3: row.append([''])
                        card_table_data.append(row)
                        row = []
                if card_table_data:
                    t = Table(card_table_data, colWidths=[5.2*cm, 5.2*cm, 5.2*cm])
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F9F9F7')),
                        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#D3D1C7')),
                        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#D3D1C7')),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('PADDING', (0,0), (-1,-1), 8),
                    ]))
                    story.append(t)

            # Form answers (perguntas orientadoras)
            form_ans = page.get('form_answers', {})
            if form_ans and any(v.strip() for v in form_ans.values()):
                story.append(Spacer(1, 0.3*cm))
                s_fq = ParagraphStyle('FQ', fontSize=10, leading=14, fontName='Helvetica-Bold', textColor=colors.HexColor('#2C2C2A'), spaceAfter=2)
                s_fa = ParagraphStyle('FA', fontSize=10, leading=14, textColor=colors.HexColor('#444441'), spaceAfter=8, leftIndent=10)
                q_n = 1
                for _grp, _qs in FORM_QUESTIONS:
                    for _q in _qs:
                        key = f'q{q_n}'
                        ans = form_ans.get(key, '').strip()
                        if ans:
                            story.append(Paragraph(f'{q_n}. {_rl_esc(_q)}', s_fq))
                            story.append(Paragraph(_rl_esc(ans), s_fa))
                        q_n += 1

            story.append(Spacer(1, 0.5*cm))

        story.append(PageBreak())

    doc.build(story)
    buf.seek(0)
    safe_name = secure_filename(project['name'].replace(' ', '_'))
    return send_file(buf, mimetype='application/pdf', as_attachment=True,
                     download_name=f'{safe_name}_estudio.pdf')

# ── Word Export ───────────────────────────────────────────────────────────────
@app.route('/api/projects/<pid>/export/docx', methods=['GET'])
@require_auth
def export_docx(pid):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3);  sec.right_margin  = Cm(3)

    # Capa
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project['name']); run.font.size = Pt(28); run.font.bold = True
    _pt_map2 = {'doc': 'Documentário', 'book': 'Livro', 'photo': 'Fotografia', 'blog': 'Blog / Não-Ficção'}
    pt = _pt_map2.get(project['type'], project['type'])
    p2 = doc.add_paragraph(pt); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(14)
    created = project.get('created', '')[:10]
    p3 = doc.add_paragraph(f'criado em {created}'); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p3.runs: p3.runs[0].font.size = Pt(10)
    doc.add_page_break()

    for section in project.get('sections', []):
        pages = section.get('pages', [])
        if not pages: continue
        doc.add_heading(section['label'].upper(), level=1)
        for page in pages:
            doc.add_heading(page.get('title') or page.get('name', ''), level=2)
            if page.get('content'):
                for line in page['content'].split('\n'):
                    line = line.strip()
                    if not line: doc.add_paragraph(''); continue
                    if line.startswith('## '):
                        doc.add_heading(line[3:], level=3)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith('> '):
                        p = doc.add_paragraph(line[2:])
                        if p.runs: p.runs[0].font.italic = True
                    else:
                        doc.add_paragraph(line)
            for scene in page.get('scenes', []):
                h = doc.add_heading(scene.get('title', 'cena'), level=3)
                if scene.get('content'):
                    doc.add_paragraph(scene['content'])
                if scene.get('notes'):
                    p = doc.add_paragraph(f'[{scene["notes"]}]')
                    if p.runs: p.runs[0].font.italic = True
            doc.add_paragraph('')
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    safe_name = secure_filename(project['name'].replace(' ', '_'))
    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,
                     download_name=f'{safe_name}_estudio.docx')

# ── Helpers ───────────────────────────────────────────────────────────────────
DOC_SECTIONS = [
    {'id': 'pesquisa', 'label': 'pesquisa', 'icon': '🔍', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'contexto histórico', 'title': 'contexto histórico', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'fontes primárias', 'title': 'fontes primárias', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'pre', 'label': 'pré-produção', 'icon': '📋', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'planejamento geral', 'title': 'planejamento geral', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': [], 'form_answers': {}},
        {'id': str(uuid.uuid4()), 'name': 'fluxo da narrativa', 'title': 'fluxo da narrativa', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': [], 'form_answers': {}},
        {'id': str(uuid.uuid4()), 'name': 'personagens', 'title': 'personagens', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': [], 'form_answers': {}},
        {'id': str(uuid.uuid4()), 'name': 'locações', 'title': 'locações', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': [], 'form_answers': {}},
        {'id': str(uuid.uuid4()), 'name': 'perguntas orientadoras', 'title': 'perguntas orientadoras', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': [], 'form_answers': {}},
    ]},
    {'id': 'producao', 'label': 'produção', 'icon': '🎥', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'plano de filmagem', 'title': 'plano de filmagem', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'entrevistas', 'title': 'entrevistas', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'cenas', 'title': 'cenas', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'pos', 'label': 'pós-produção', 'icon': '✂️', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'edição', 'title': 'edição', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'trilha sonora', 'title': 'trilha sonora', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'distribuicao', 'label': 'distribuição', 'icon': '🌐', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'festivais', 'title': 'festivais', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'plataformas', 'title': 'plataformas', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
]

BOOK_SECTIONS = [
    {'id': 'pesquisa', 'label': 'pesquisa', 'icon': '🔍', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'notas de campo', 'title': 'notas de campo', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'contexto', 'title': 'contexto', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'materiais', 'label': 'materiais', 'icon': '📎', 'isUpload': True, 'files': [], 'pages': []},
    {'id': 'roteiro', 'label': 'roteiro', 'icon': '📝', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'estrutura geral', 'title': 'estrutura geral', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'arco narrativo', 'title': 'arco narrativo', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'personagens', 'label': 'personagens', 'icon': '👤', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'protagonista', 'title': 'protagonista', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'cenarios', 'label': 'cenários', 'icon': '🗺️', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'cenário principal', 'title': 'cenário principal', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'capitulos', 'label': 'capítulos', 'icon': '📑', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'capítulo 1', 'title': 'capítulo 1', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'capítulo 2', 'title': 'capítulo 2', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'narrativa', 'label': 'narrativa', 'icon': '✍️', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'voz narrativa', 'title': 'voz narrativa', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'estilo', 'title': 'estilo', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
]

PHOTO_SECTIONS = [
    {'id': 'albuns', 'label': 'álbuns', 'icon': '📷', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'álbum 1', 'title': 'álbum 1', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'ensaios', 'label': 'ensaios', 'icon': '🎞️', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'ensaio 1', 'title': 'ensaio 1', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'referencias', 'label': 'referências visuais', 'icon': '🖼️', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'moodboard', 'title': 'moodboard', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'publicacoes', 'label': 'publicações', 'icon': '🌐', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'instagram', 'title': 'instagram', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
]

BLOG_SECTIONS = [
    {'id': 'artigos', 'label': 'artigos', 'icon': '✍️', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'artigo 1', 'title': 'artigo 1', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'pesquisa', 'label': 'pesquisa', 'icon': '🔍', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'fontes', 'title': 'fontes', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
        {'id': str(uuid.uuid4()), 'name': 'notas de campo', 'title': 'notas de campo', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'rascunhos', 'label': 'rascunhos', 'icon': '📝', 'pages': [
        {'id': str(uuid.uuid4()), 'name': 'rascunho 1', 'title': 'rascunho 1', 'content': '', 'scenes': [], 'images': [], 'board_cards': [], 'board_connections': []},
    ]},
    {'id': 'publicados', 'label': 'publicados', 'icon': '✅', 'pages': []},
]

def build_sections(ptype):
    import copy
    templates = {'doc': DOC_SECTIONS, 'book': BOOK_SECTIONS, 'photo': PHOTO_SECTIONS, 'blog': BLOG_SECTIONS}
    template = templates.get(ptype, BLOG_SECTIONS)
    sections = copy.deepcopy(template)
    for s in sections:
        s['id'] = s.get('id', str(uuid.uuid4()))
        for pg in s.get('pages', []):
            pg['id'] = str(uuid.uuid4())
    return sections

FORM_QUESTIONS = [
    ('Visão e Problema', [
        'What is the core problem you are trying to help solve?',
        'What needs are we trying to address?',
        'Who is the intended audience we hope to impact?',
        'Why is it important?',
        'What is the purpose of our film\'s engagement/impact campaign?',
        'What specific insights and learnings might emerge from this work that could inform existing social movements?',
    ]),
    ('Teoria da Mudança', [
        'What factors contribute to the issue/problem identified in our vision statement?',
        'What needs to change for our goal to be realized?',
        'What are the processes for change?',
        'How can our project influence or contribute to that process? Who controls whether that change happens?',
        'What needs to happen for decision makers to take action?',
        'Why are they not taking action now?',
        'Who benefits from the status quo?',
        'How do they benefit?',
        'What level of power do they have?',
        'Who is adversely affected by the status quo? How?',
        'What level of power do they have?',
    ]),
    ('Stakeholders e Parcerias', [
        'Which stakeholders are working on this issue already?',
        'What strategies are they employing?',
        'Are these strategies working or not?',
        'What role(s) can our project play in this environment?',
        'What campaign approaches and activities could accomplish this?',
        'What short-term outcomes will these activities achieve?',
        'What are our most important priorities?',
        'Given the stakeholders we have identified, which are our film naturally aligned with?',
        'Do the stakeholders collaborate, and if so, how? If not, why?',
        'Which stakeholders should we prioritize based on their reach, influence, experience or reputation?',
        'Does the work of these stakeholders align or overlap with our change strategy?',
        'For each stakeholder, how is the organization structured and how are partnership decisions made?',
        'Do we already have relationships with these organizations? If not, through which channels might we approach them?',
        'How might we describe our film and social change goals to these partners to pique their interest?',
        'How will we build trust with each partner and develop shared goals?',
        'What kind of communication will we need to nurture the relationship over time?',
        'How might we structure these partnerships to ensure that they are mutually beneficial and sustainable?',
        'Do we see these partnerships as long-term or short-term?',
    ]),
    ('Engajamento', [
        'When thinking about a ladder of engagement for each target audience, what is our call to action for each that will add value to the movement?',
        'When the lights go up after our film screens and someone asks what he or she can do to help, how will we answer?',
        'What kinds of engagement activities will we use to realize those calls to action? (e.g. viewers guides, curriculum, a "take action" website, a petition, a social media conversation)',
    ]),
    ('Equipe', [
        'What are the assets of our current team — including skills, availability, experience, interest and commitment?',
        'Does our team have extensive project management experience?',
        'What are the gaps in our current team — including skills, availability, experience, interest and commitment?',
        'What kinds of skills do we need to add to our team in order to meet the needs of our campaign?',
        'How would we prioritize those?',
        'What are the different ways that our team could be structured?',
        'What are the pros and cons of each type of structure?',
        'Are there ways that we could deepen our partnerships to meet some of our staffing needs?',
        'What does our personnel budget look like?',
    ]),
    ('Cronograma', [
        'When will we begin our campaign?',
        'Are there broad phases within our campaign (i.e. broadcast phase, community screening phase, etc.)?',
        'What are the key dates that affect our campaign? Think about holidays, remembrance months, political cycles, distribution opportunities, etc.',
        'What key milestones should we focus on to ensure we stay on track?',
        'What are the fundraising deadlines and decision dates that we have targeted?',
        'What are the important dates for our partners?',
        'When do we envision wrapping up our campaign and what is our exit strategy?',
    ]),
    ('Orçamento', [
        'What are the key activities/components related to our campaign?',
        'What are key costs associated with each activity?',
        'What drives each cost? Staff? Consultants? Events? Travel? Social Media? Web Design? Screenings? Data Collection?',
        'Based on our timeline, how would we prioritize the activities?',
    ]),
    ('Captação de Recursos', [
        'What are our potential sources of revenue?',
        'Do you have a clear understanding of what kind of support you need?',
        'Have you identified funders who support these issues?',
        'Are there funder networks that focus on these issues? Which funders supported your film\'s production? Would they support the impact campaign?',
        'What are the funding guidelines and criteria for each of the potential funders you have identified?',
        'What level of support does each typically give? When are their deadlines and what is the application process?',
        'Have you thought about the best way to approach different funders?',
        'Do you have personal relationships with key staff? Is it possible to initiate a conversation before submitting a proposal?',
        'If our film is nominated for, or wins, an award, how will we leverage that opportunity?',
        'If our film or campaign receives a major newspaper endorsement or mention, how will we take advantage of that opportunity?',
        'If an influencer or celebrity talks about our film, how will we use that opportunity?',
        'If the issue(s) in our film begin receiving major attention from elected officials and there is political movement, how will we ensure our film becomes part of that conversation?',
    ]),
    ('Riscos e Contingências', [
        'What if one of our campaign partners or funders does not like or approve of the final cut of the film?',
        'What will we do if one of our partner organizations receives bad press?',
        'What are other "unexpected" things that we should be thinking about?',
    ]),
]

def _rl_esc(text):
    """Escape special XML chars for ReportLab Paragraph."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

# ── Form PDF Export ───────────────────────────────────────────────────────────
@app.route('/api/projects/<pid>/export/form/<pgid>', methods=['GET'])
@require_auth
def export_form_pdf(pid, pgid):
    data = load_data()
    project = next((p for p in data['projects'] if p['id'] == pid), None)
    if not project: abort(404)
    page = None
    for s in project.get('sections', []):
        for pg in s.get('pages', []):
            if pg['id'] == pgid:
                page = pg
                break
    if not page: abort(404)

    answers = page.get('form_answers', {})
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm)

    TYPE_COLOR = colors.HexColor('#3C3489')
    s_cover = ParagraphStyle('CT', fontSize=26, leading=32, textColor=TYPE_COLOR, spaceAfter=6, fontName='Helvetica-Bold')
    s_proj = ParagraphStyle('CP', fontSize=14, textColor=colors.HexColor('#888780'), spaceAfter=4)
    s_meta = ParagraphStyle('CM', fontSize=10, textColor=colors.HexColor('#B4B2A9'))
    s_grp = ParagraphStyle('GH', fontSize=13, leading=18, fontName='Helvetica-Bold', textColor=TYPE_COLOR, spaceBefore=18, spaceAfter=4)
    s_q = ParagraphStyle('Q', fontSize=11, leading=16, textColor=colors.HexColor('#2C2C2A'), spaceAfter=3, fontName='Helvetica-Bold')
    s_a = ParagraphStyle('A', fontSize=11, leading=17, textColor=colors.HexColor('#444441'), spaceAfter=10, leftIndent=12)
    s_empty = ParagraphStyle('AE', fontSize=11, leading=17, textColor=colors.HexColor('#BBBAB2'), spaceAfter=10, leftIndent=12)

    story = []
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(page.get('title', 'Perguntas Orientadoras'), s_cover))
    story.append(Paragraph(project['name'], s_proj))
    story.append(Paragraph(f'exportado em {datetime.now().strftime("%d/%m/%Y")}', s_meta))
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width='100%', thickness=2, color=TYPE_COLOR))
    story.append(PageBreak())

    q_num = 1
    for group_name, questions in FORM_QUESTIONS:
        story.append(Paragraph(group_name.upper(), s_grp))
        story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#D3D1C7')))
        story.append(Spacer(1, 0.15*cm))
        for q in questions:
            key = f'q{q_num}'
            ans = answers.get(key, '').strip()
            story.append(Paragraph(f'{q_num}. {_rl_esc(q)}', s_q))
            if ans:
                for line in ans.split('\n'):
                    line = line.strip()
                    if line:
                        story.append(Paragraph(_rl_esc(line), s_a))
                    else:
                        story.append(Spacer(1, 0.1*cm))
            else:
                story.append(Paragraph('—', s_empty))
            q_num += 1
        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    buf.seek(0)
    safe_name = secure_filename(f'{project["name"]}_perguntas_orientadoras'.replace(' ', '_'))
    return send_file(buf, mimetype='application/pdf', as_attachment=True,
                     download_name=f'{safe_name}.pdf')


# ── Data import (one-time restore) ───────────────────────────────────────────
@app.route('/api/import', methods=['GET'])
@require_auth
def import_page():
    """Simple upload form to restore projects.json from local machine."""
    return '''<!DOCTYPE html><html><head><meta charset="UTF-8">
    <title>importar dados</title>
    <style>body{font-family:monospace;background:#EDE8E0;display:flex;align-items:center;justify-content:center;height:100vh;margin:0}
    .box{background:#F7F4EF;border:0.5px solid #D8D2C8;border-radius:12px;padding:40px;text-align:center;max-width:400px}
    h2{margin:0 0 8px;font-size:16px}p{color:#5C6470;font-size:13px;margin:0 0 24px}
    input{display:block;margin:0 auto 16px;font-family:monospace}
    button{background:#1A2744;color:#fff;border:none;padding:10px 24px;border-radius:8px;cursor:pointer;font-family:monospace;font-size:13px}
    .warn{color:#A31C2C;font-size:12px;margin-top:12px}</style></head>
    <body><div class="box"><h2>importar projects.json</h2>
    <p>Selecione o arquivo projects.json do seu computador para restaurar os dados no servidor.</p>
    <form method="POST" enctype="multipart/form-data">
      <input type="file" name="file" accept=".json" required />
      <button type="submit">restaurar dados</button>
    </form>
    <p class="warn">⚠ isso sobrescreve todos os dados existentes no servidor</p>
    </div></body></html>'''

@app.route('/api/import', methods=['POST'])
@require_auth
def import_data():
    """Restore projects.json uploaded from local machine."""
    if 'file' not in request.files:
        return 'nenhum arquivo enviado', 400
    f = request.files['file']
    try:
        data = json.loads(f.read().decode('utf-8'))
        if 'projects' not in data:
            return 'arquivo inválido — falta chave "projects"', 400
        save_data(data)
        count = len(data['projects'])
        return f'''<!DOCTYPE html><html><head><meta charset="UTF-8">
        <title>importado</title>
        <style>body{{font-family:monospace;background:#EDE8E0;display:flex;align-items:center;justify-content:center;height:100vh;margin:0}}
        .box{{background:#F7F4EF;border:0.5px solid #D8D2C8;border-radius:12px;padding:40px;text-align:center}}
        h2{{color:#2B5231}}a{{color:#1A2744}}</style></head>
        <body><div class="box"><h2>✓ {count} projeto(s) restaurado(s)</h2>
        <p><a href="/">← voltar ao estúdio</a></p></div></body></html>'''
    except Exception as e:
        return f'erro ao processar arquivo: {e}', 400

# ── Backup status / manual trigger ───────────────────────────────────────────
@app.route('/api/backup', methods=['POST'])
@require_auth
def trigger_backup():
    """Force an immediate backup of projects.json to Google Drive."""
    if not GDRIVE_FOLDER_ID or not os.environ.get('GDRIVE_SERVICE_ACCOUNT_B64'):
        return jsonify({'status': 'disabled', 'message': 'Google Drive não configurado'}), 200
    _backup_file_to_drive(DATA_FILE, 'projects.json')
    return jsonify({'status': 'ok', 'message': 'backup enviado para o Google Drive'})

@app.route('/api/backup/status', methods=['GET'])
@require_auth
def backup_status():
    """Check if Drive backup is configured."""
    configured = bool(GDRIVE_FOLDER_ID and os.environ.get('GDRIVE_SERVICE_ACCOUNT_B64'))
    return jsonify({'configured': configured, 'folder_id': GDRIVE_FOLDER_ID or None})


# ── Startup ───────────────────────────────────────────────────────────────────
# Restore projects.json from Drive if disk is empty (e.g. first deploy or disk reset)
with app.app_context():
    _restore_from_drive()

if __name__ == '__main__':
    app.run(debug=True, port=5000)
